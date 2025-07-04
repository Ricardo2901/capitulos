"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""
def capitulo6():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 6
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo VI.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.-	Análisis comparativo de la composición florística y faunística del área sujeta a Cambio de uso de suelo en Terrenos Forestales con relación a los tipos de vegetación del ecosistema de la cuenca, subcuenca o microcuenca hidrográfica, que permita determinar el grado de afectación por el Cambio de Uso de Suelo en Terrenos forestales.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.- Comparativos de composición de la vegetación presente dentro del ACUSTF y Sistema Ambiental.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('En este capítulo se analizarán _______________________________________________________________ con sus cuatro comparativos en composición florística: Comparativo de individuos e índice de valor de importancia; es decir, total de individuos extrapolados y su porcentaje de afectación al Sistema Ambiental por el Cambio y uso de Suelo, Comparativo por índices de biodiversidad, Comparativo por valor densidad de especies y Comparativos de Índices de similitud/disimilitud.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.1- Comparativo de individuos e Índice de Valor de Importancia por estrato del Sistema Ambiental –ACUSTF en el ______________________')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Según Aguirre (1999) el índice de valor de importancia (IVI), indica que tan importante es una especie dentro de la comunidad. Las especies que tienen el IVI más alto significa entre otras cosas que es dominante ecológicamente: que absorbe muchos nutrientes, que ocupa mayor espacio físico, que controla en un porcentaje alto la energía que llega a este sistema. Este índice sirve para comparar el peso ecológico de cada especie dentro del ecosistema. Para calcular este parámetro se utiliza la Densidad relativa, Frecuencia relativa y Dominancia relativa.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('A continuación, se describen las fórmulas que se utilizaron para la estimación del Índice de Valor de Importancia.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDensidad relativa. Está dada por el resultado de la densidad absoluta entre el número total de todos los individuos muestreados expresados en porcentajes ')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Formula del Capitulo 6.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_2.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del Capitulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDer = Densidad Relativa'
                                    '\nNi = Número de individuos de la especie'
                                    '\nNt = Número total de individuos de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nFrecuencia relativa. Es el resultado de dividir la frecuencia absoluta de cada especie entre el número total de esas especies expresadas en porcentajes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_6.png', width=Cm(4.74), height=Cm(1.50))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nFr = Frecuencia relativa'
                                    '\nFai = Frecuencia absoluta de cada especie'
                                    '\nFat = Frecuencia absoluta de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nLa dominancia relativa. Se calcula como la proporción de una especie en el área total evaluada, expresada en porcentaje.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_4.png', width=Cm(4.79), height=Cm(1.50))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDor = Densidad relativa'
                                    '\nDai = Densidad absoluta de una especie'
                                    '\nDat= Densidad absoluta total de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de valor de importancia (IVI). El índice de valor de importancia define cuáles de las especies presentes contribuyen en el carácter y estructura de una Comunidad. Este valor se obtiene mediante la sumatoria de la frecuencia relativa, la densidad relativa y la dominancia relativa.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_7.png', width=Cm(4.99), height=Cm(1.20))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nIVI = Índice de Valor de Importancia'
                                    '\nDer = Densidad relativa'
                                    '\nDor = Dominancia relativa'
                                    '\nFr = Frecuencia relativa')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capítulo 6.1.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 6.1.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.1.- Comparativo por total de individuos e Índice de Valor de Importancia en el ____')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.1 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescribir los del cuadro.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver  resto del contenido del capítulo 6.1.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nEn general las especies de lento crecimiento y las enlistadas en la NOM- 059- SEMARNAT 2010 se rescatarán y reubicarán a una superficie que tenga las mismas condiciones donde se distribuyen actualmente para que no pierdan su germoplasma.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina en el capitulo 6.1.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 6.1.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'VI.1.2.- Comparativo por índices de biodiversidad del Sistema Ambiental –ACUSTF del _____________________________.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices basados en la dominancia son parámetros inversos al concepto de uniformidad o equidad de la comunidad. Toman en cuenta la representatividad de las especies con mayor valor de importancia sin evaluar la contribución del resto de las especies. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para medir la dominancia de las especies los índices de biodiversidad más comunes son: Simpson y Berger Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 6.1.2 ###
    #########################
    listaValores612 = [
        '0 - 0.33',
        '0.34 - 0.66',
        '> 0.67',
    ]

    diversidadSignificancia612 = [
        'Diversidad Baja',
        'Diversidad Media',
        'Diversidad Alta',
    ]

    heterogeneoHomogeneoSignificancia612 = [
        'Heterogéneo en abundancia',
        'Ligeramente Heterogéneo en en abundancia',
        'Homogéneo en abundancia',
    ]

    valores612 = range(len(listaValores612))
    diversidad612 = range(len(diversidadSignificancia612))
    heterogeneoHomogeneo612 = range(len(heterogeneoHomogeneoSignificancia612))

    filasCap612 = len(valores612) + 2
    
    tabla6 = doc.add_table(rows=filasCap612, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(3.46)
        rows.cells[1].width = Cm(3.66)
        rows.cells[2].width = Cm(9.66)

    #########################
    # Celda fusionada "Escalas de interpretación de significancia 0-1"
    row1 = tabla6.rows[0]
    merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Escalas de interpretación de significancia 0-1')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell1, '0070C0')

    #########################
    # Celda fusionada "Significancia"
    row1 = tabla6.rows[1]
    merged_cell1 = row1.cells[1].merge(row1.cells[1].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Significancia')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla6.cell(1, 0)
    t6 = cell.paragraphs[0].add_run('Valores')
    t6.font.size = Pt(12)
    t6.font.name = 'Arial'
    t6.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in valores612:
        cell = tabla6.cell(cols + 2, 0)
        t6 = cell.paragraphs[0].add_run(f'{listaValores612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        t6.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in diversidad612:
        cell = tabla6.cell(cols + 2, 1)
        t6 = cell.paragraphs[0].add_run(f'{diversidadSignificancia612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in heterogeneoHomogeneo612:
        cell = tabla6.cell(cols + 2, 2)
        t6 = cell.paragraphs[0].add_run(f'{heterogeneoHomogeneoSignificancia612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el Índice de Margalef el criterio es de 2-5, donde, sus escalas de interpretación son: de 0-2 se considera diversidad baja, de 2-5 se considera diversidad media y mayor de 5 se considera diversidad alta y el Índice se Shannon tiene un criterio de 2-3 donde su escala de interpretación es: 0-2 se considera diversidad baja, de 2-3 se considera diversidad media y mayor de 3 se considera diversidad alta. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDe acuerdo al análisis realizado en el área de cambio de uso de suelo y sistema ambiental se tiene lo siguiente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.2.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.2.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.2.1.- Riqueza específica') 
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('La riqueza específica (S) es la forma más sencilla de medir la biodiversidad, ya que se basa únicamente en el número de especies presentes, sin tomar en cuenta el valor de importancia de las mismas. La forma ideal de medir la riqueza específica es contar con un inventario completo que nos permita conocer el número total de especies (S) obtenido por un censo de la comunidad. Esto es posible únicamente para ciertas taxas bien conocidos y de manera puntual en tiempo y en espacio. La mayoría de las veces tenemos que recurrir a índices de riqueza específica obtenidos a partir de un muestreo de la comunidad. A continuación, se describen los índices más comunes para medir la riqueza de especies de acuerdo a (Moreno 2001)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una Comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDmg = Índice de Margalef'
                                    '\nS = Número de especies.'
                                    '\nN = Número total de individuos'
                                    '\nD = Densidad'
                                    '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                    '\nLn= Logaritmo natural')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.2.- Riqueza de especies (Índice de Margalef)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.1.- Riqueza de especies (Índice de Margalef)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De acuerdo al cuadro y gráfico anterior se observa que, en cuanto a Riqueza de especies para los estratos, arbustivo y suculento presentan un valor medio para las dos áreas (ACUSTF y Sistema Ambiental) para el estrato gramíneo ambas áreas presentan valores bajos y para el estrato herbáceo en el ACUSTF valor medio mientras que en el SA el valor es bajo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.1 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.3.- Riqueza de especies (Índice de Menhinick)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.2.- Riqueza de especies (Índice de Menhinick)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.1.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.2.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.2.2.- Dominancia de especies ')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices basados en la dominancia son parámetros inversos al concepto de uniformidad o equidad de la comunidad. Toman en cuenta la representatividad de las especies con mayor valor de importancia sin evaluar la contribución del resto de las especies. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para medir la dominancia de las especies los índices de biodiversidad más comunes son: Simpson y Berger Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Índice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. Manifiesta la probabilidad de que dos individuos tomados al azar de una muestra sean de la misma especie. Está fuertemente influido por la importancia de las especies más dominantes (Magurran, 1988; Peet, 1974). Es decir, cuanto más se acerca el valor de este índice a la unidad existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.1.2.2 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6122/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde,'
                                        '\nƛ = índice de dominancia se Simpson'
                                        '\nID=índice de diversidad'
                                        '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra'
    )
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia, de 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia y mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.4.- Dominancia de especies (Índice de Simpson)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.3.- Dominancia de especies (Índice de Simpson)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('FAVOR DE DESCRIBIR EL RESTO DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia (Magurran, 1988).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.2 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.5.- Dominancia de especies (Berger Parker)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.4.- Dominancia de especies (Berger Parker)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.3.- Comparativo por valor densidad de especies en el Sistema Ambiental -ACUSTF en el ______________________________.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para realizar este comparativo se extrapolaron los individuos de las especies a hectáreas, y se calificó de acuerdo a los siguientes cuadros que mencionan los valores de densidad. Es decir, la densidad de individuos por hectáreas y su respectiva calificación si es vegetación Rala, Semidensa y Densa. Estos cuadros fueron extraídos de la Guía de Métodos para medir la biodiversidad de la revista Área Agropecuaria y de Recursos Naturales Renovables de Ecuador.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.5.- Dominancia de especies (Berger Parker)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.3 ###
    #########################
    # Densidad de los arboles #
    densidadVegetacion = [
        'los árboles', 'los arbustos', 'las hierbas'
    ]

    densidadArboles = [
        '0-300 Individuos/hectárea',
        '301-600 Individuos/hectárea',
        'más de 600 Individuos/hectárea',
    ]

    densidadArbustos = [
        '0-500 Individuos/hectárea',
        '501-1000 Individuos/hectárea',
        'más de 1000 Individuos/hectárea',
    ]

    densidadHierbas = [
        '0-1000 Individuos/hectárea',
        '1001-2000 Individuos/hectárea',
        'más de 2000 Individuos/hectárea',
    ]

    valorPonderado = [
        1.67,
        3.33,
        5,
    ]

    calificacionCap613 = [
        'Vegetación Rala (R)',
        'Vegetación Semidensa (SD)',
        'Vegetación Densa (D)',
    ]

    columnasCap613 = [
        'Valor Calculado de Densidad',
        'Valor Ponderado',
        'Clasificación',
    ]

    densidadVegetacionRango = range(len(densidadVegetacion))
    densidadArbolesRango = range(len(densidadArboles))
    densidadArbustosRango = range(len(densidadArbustos))
    valorPonderadoRango = range(len(valorPonderado))
    densidadHierbasRango = range(len(densidadHierbas))
    calificacionCap613Rango = range(len(calificacionCap613))
    columnasCap613Rango = range(len(columnasCap613))

    VegetacionRango = len(densidadVegetacion)
    ArbolesRango = len(densidadArboles)
    ArbustosRango = len(densidadArbustos)
    ponderadoRango = len(valorPonderado)
    HierbasRango = len(densidadHierbas)
    cap613Rango = len(calificacionCap613)
    cap613ColumnasRango = len(columnasCap613)
    
    tabla6 = doc.add_table(rows=15, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(7.77)
        rows.cells[1].width = Cm(4.06)
        rows.cells[2].width = Cm(6.62)

    #########################
    # Celdas fusionadas "Valores de densidad para estimar la densidad de ..."
    for celda_fusionada in densidadVegetacionRango:
        i = celda_fusionada * 5

        row1 = tabla6.rows[i]
        merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

        # Agregar texto a la celda fusionada
        t6 = merged_cell1.paragraphs[0].add_run(f'Valores de densidad para estimar la densidad de {densidadVegetacion[celda_fusionada]}')
        t6.font.name = 'Arial'
        t6.font.size = Pt(12)
        t6.bold = True
        merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(merged_cell1, '4F81BD')

    for rows in densidadVegetacionRango:
        i = (rows * 5) + 1

        for cols in densidadVegetacionRango:
            cell = tabla6.cell(i, cols)
            t6 = cell.paragraphs[0].add_run(f'{columnasCap613[cols]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            t6.bold = True

    for cols in densidadVegetacionRango:
        i = (cols * 5) + 2

        for valor in densidadVegetacionRango:
            k = i + valor
            cell = tabla6.cell(k, 1)
            t6 = cell.paragraphs[0].add_run(f'{valorPonderado[valor]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for calif in densidadVegetacionRango:
            k = i + calif
            cell = tabla6.cell(k, 2)
            t6 = cell.paragraphs[0].add_run(f'{calificacionCap613[calif]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for densidad in range(1):
            for arboles in range(1):
                k = i + densidad
                cell = tabla6.cell(k, 0)
                t6 = cell.paragraphs[0].add_run(f'{densidadArboles[arboles]}')
                t6.font.size = Pt(12)
                t6.font.name = 'Arial'
                tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for arbusto in range(1):
                k = i + densidad
                cell = tabla6.cell(k, 0)
                t6 = cell.paragraphs[0].add_run(f'{densidadArbustos[arbusto]}')
                t6.font.size = Pt(12)
                t6.font.name = 'Arial'
                tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for hierbas in range(1):
                k = i + densidad
                cell = tabla6.cell(k, 0)
                t6 = cell.paragraphs[0].add_run(f'{densidadHierbas[hierbas]}')
                t6.font.size = Pt(12)
                t6.font.name = 'Arial'
                tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 6 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo6() # Crear el documento